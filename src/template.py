import pandas as pd
from connect_db import getConnection
import json
from docx import Document
from io import BytesIO
from datetime import datetime
import re
import os
import sys
import subprocess
import pikepdf
import traceback
import sys
import psycopg2

def getCatergoryDropDownData():
    conn = getConnection()
    sql = """
    SELECT LOWER(REPLACE(ct2.name, ' ', '_')) as casetypevalue, 
           ct2.name as casetypelabel,
           LOWER(REPLACE(templateName, ' ', '_')) as value, 
           templateName as label, 
           ct.price
    FROM template.TemplateTypes ct 
    INNER JOIN template.CaseTypes ct2 ON ct.caseTypeid = ct2.id 
    WHERE ct.isActive = True
    """
    df = pd.read_sql_query(sql, conn)
    grouped_dict = df.groupby("casetypevalue").apply(
        lambda g: g[["value", "label"]].to_dict(orient="records")
    ).to_dict()
    case_types_array = df[["casetypevalue", "casetypelabel"]].drop_duplicates().rename(
        columns={"casetypevalue": "value", "casetypelabel": "label"}
    ).to_dict(orient="records")
    data = {
        "CaseTypes": case_types_array,
        "TemplateTypes": grouped_dict,
        'rawData': df.to_dict(orient="records")
    }
    return data


def getDatafromTable(tableName):
    conn = getConnection()
    sql = f'SELECT * FROM template.{tableName}'
    df = pd.read_sql_query(sql, conn)
    df = df.to_dict(orient='records')
    return df


def getDataSet():
    conn = getConnection()
    sql = """
    SELECT 'CaseTypes' as type, id, LOWER(REPLACE(name, ' ', '_')) AS value, name AS label
    FROM template.CaseTypes
    UNION ALL
    SELECT 'TemplateTypes' as type, id, LOWER(REPLACE(templateName, ' ', '_')) AS value, templateName AS label
    FROM template.TemplateTypes;
    """
    df = pd.read_sql_query(sql, conn)
    df = df.to_dict(orient='records')
    case_types = [item for item in df if item['type'] == 'CaseTypes']
    template_types = [item for item in df if item['type'] == 'TemplateTypes']
    return {
        "CaseTypes": case_types,
        "TemplateTypes": template_types
    }


def updateRecordInTable(table_name, records):
    conn = getConnection()
    cursor = conn.cursor()

    result = {
        'updated': 0,
        'added': 0
    }

    full_table = "template." + table_name

    for record in records:
        record_id = record.get('id')

        if record.get('updated'):
            if not record_id:
                print("Missing 'id' for update operation.")
                continue

            cursor.execute(f'SELECT 1 FROM {full_table} WHERE id=%s', (record_id,))
            exists = cursor.fetchone()

            update_data = {k: v for k, v in record.items() if k not in ['id', 'updated', 'added']}
            if not update_data:
                print(f"No fields to update or insert for id {record_id}.")
                continue

            if exists:
                set_clause = ', '.join(f'"{k}"=%s' for k in update_data)
                values = list(update_data.values()) + [record_id]

                query = f'UPDATE {full_table} SET {set_clause} WHERE id=%s'
                cursor.execute(query, values)
                result['updated'] += 1
            else:
                insert_data = update_data.copy()
                insert_data['id'] = record_id

                columns = ', '.join(f'"{k}"' for k in insert_data)
                placeholders = ', '.join(['%s'] * len(insert_data))
                values = list(insert_data.values())

                query = f'INSERT INTO {full_table} ({columns}) VALUES ({placeholders})'
                cursor.execute(query, values)
                result['added'] += 1

        elif record.get('added'):
            insert_data = {k: v for k, v in record.items() if k not in ['updated', 'added']}
            if not insert_data:
                print("No fields to insert.")
                continue

            columns = ', '.join(f'"{k}"' for k in insert_data)
            placeholders = ', '.join(['%s'] * len(insert_data))
            values = list(insert_data.values())

            query = f'INSERT INTO {full_table} ({columns}) VALUES ({placeholders})'
            cursor.execute(query, values)
            result['added'] += 1

    conn.commit()
    conn.close()
    return result


def getTemplateFeilds(caseType, TemplateType):
    conn = getConnection()
    sql = f"""
    SELECT templateForm 
    FROM template.Templates t 
    WHERE t.caseTypeid = (SELECT id FROM template.CaseTypes WHERE name = %s) 
    AND t.TemplateTypeid = (SELECT id FROM template.TemplateTypes tt WHERE templateName = %s)
    """
    df = pd.read_sql_query(sql, conn, params=(caseType, TemplateType))
    if not df.empty:
        df = df.iloc[0][0]
        df = df.replace("'", '"')
        df = json.loads(df)
        return df
    else:
        return {}


def updateTemplateFields(caseType, TemplateType, updatedData):
    conn = getConnection()
    cursor = conn.cursor()
    updated_json = json.dumps(updatedData).replace('"', "'")
    sql = f"""
    UPDATE template.Templates
    SET templateForm = %s
    WHERE caseTypeid = (SELECT id FROM template.CaseTypes WHERE name = %s)
    AND TemplateTypeid = (SELECT id FROM template.TemplateTypes WHERE templateName = %s)
    """
    try:
        cursor.execute(sql, (updated_json, caseType, TemplateType))
        conn.commit()
        return "Template updated successfully."
    except Exception as e:
        return "Error updating template:" + str(e)
    finally:
        conn.close()


def extractDataItems(document, caseType, TemplateType):
    doc_bytes = document.read()
    doc = Document(BytesIO(doc_bytes))
    full_text = ''
    for para in doc.paragraphs:
        full_text += para.text + '\n'
    placeholders = re.findall(r'\[([^\[\]]+)\]', full_text)
    placeholdersData = {item: item for item in placeholders}
    conn = getConnection()
    cursor = conn.cursor()
    try:
        cursor.execute('SELECT id FROM template.CaseTypes WHERE name = %s', (caseType,))
        case_row = cursor.fetchone()
        if not case_row:
            raise ValueError(f"CaseType '{caseType}' not found.")
        caseTypeid = case_row[0]
        cursor.execute('SELECT id FROM template.TemplateTypes WHERE templateName = %s', (TemplateType,))
        template_row = cursor.fetchone()
        TemplateTypeid = template_row[0]
        cursor.execute(
            '''
            SELECT id FROM template.Templates
            WHERE caseTypeid = %s AND TemplateTypeid = %s
            ''',
            (caseTypeid, TemplateTypeid)
        )
        existing = cursor.fetchone()
        if existing:
            cursor.execute(
                '''
                UPDATE template.Templates 
                SET templateData = %s, templateForm = %s
                WHERE caseTypeid = %s AND TemplateTypeid = %s
                ''',
                (psycopg2.Binary(doc_bytes), str(placeholdersData), caseTypeid, TemplateTypeid)
            )
            print("Existing template updated.")
        else:
            cursor.execute(
                '''
                INSERT INTO template.Templates(caseTypeid, TemplateTypeid, templateData, templateForm)
                VALUES (%s, %s, %s, %s)
                ''',
                (caseTypeid, TemplateTypeid, psycopg2.Binary(doc_bytes), str(placeholdersData))
            )
            print("New template inserted.")
        conn.commit()
    except Exception as e:
        print("Error occurred:", e)
        traceback.print_exc()  #
        conn.rollback()
        print("Error in extractDataItems:", e)
    finally:
        conn.close()
    return placeholdersData


def replace_placeholders_in_paragraph(paragraph, replacements):
    full_text = ''.join(run.text for run in paragraph.runs)
    for key, value in replacements.items():
        pattern = rf"\[{re.escape(key)}\]"
        full_text = re.sub(pattern, value, full_text)
    for run in paragraph.runs:
        run.text = ''
    if paragraph.runs:
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)


def generateProtectedPDF(caseType, TemplateType, replacements):
    conn = getConnection()
    cursor = conn.cursor()
    sql = f"""
    SELECT templateData 
    FROM template.Templates t 
    WHERE t.caseTypeid = (SELECT id FROM template.CaseTypes WHERE name = %s) 
    AND t.TemplateTypeid = (SELECT id FROM template.TemplateTypes tt WHERE templateName = %s)
    """
    cursor.execute(sql, (caseType, TemplateType))
    row = cursor.fetchone()
    conn.close()
    doc_blob = row[0]
    doc_stream = BytesIO(doc_blob)
    doc = Document(doc_stream)
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, replacements)
    today_date = datetime.now().strftime("%Y%m%d%H%M%S%f")[:-3]
    mainPath = os.path.join(os.getcwd(), 'temp')
    if not os.path.exists(mainPath):
        os.mkdir(mainPath)
    temp_docx = os.path.join(mainPath, f"{today_date}{TemplateType.replace(' ', '_')}.docx")
    doc.save(temp_docx)
    output_pdf = os.path.join(mainPath, f"{today_date}{TemplateType.replace(' ', '_')}.pdf")
    if sys.platform == "win32":
        import pythoncom
        from docx2pdf import convert
        pythoncom.CoInitialize()
        convert(temp_docx, output_pdf)
    else:
        unoconv_path = '/usr/bin/unoconv'
        # subprocess.run([unoconv_path, '-f', 'pdf', '-o', output_pdf, temp_docx])
        subprocess.run(
            [unoconv_path, '-f', 'pdf', '-o', output_pdf, temp_docx],
            env={"PATH": "/usr/bin:/bin"},  # minimal clean env
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=True
        )
    with pikepdf.open(output_pdf) as pdf:
        encrypted_pdf_stream = BytesIO()
        pdf.save(
            encrypted_pdf_stream,
            encryption=pikepdf.Encryption(
                user="",
                owner="secret",
                allow=pikepdf.Permissions(
                    extract=False,
                    modify_annotation=False,
                    modify_assembly=False,
                    modify_form=False,
                    modify_other=False,
                    print_lowres=True,
                    print_highres=True
                )
            )
        )
        return encrypted_pdf_stream, today_date


def checkUserBalance(id, caseType, templateType, replacements):
    conn = getConnection()
    cursor = conn.cursor()
    cursor.execute(f"""
    SELECT ct.price 
    FROM template.TemplateTypes ct
    WHERE ct.caseTypeid = (SELECT id FROM template.CaseTypes WHERE name = '{caseType}') 
    AND ct.templatename  =  '{templateType}'
    """)
    result = cursor.fetchone()
    if not result:
        return False, {'message': 'Template not found or inactive'}, 404
    price = result[0]
    # Check user balance
    cursor.execute("SELECT credits FROM template.users WHERE id = %s", (id,))
    user_balance = cursor.fetchone()
    if not user_balance or user_balance[0] < price:
        return False, {'message': 'Insufficient balance'}, 403
    # Deduct price from user balance
    cursor.execute("UPDATE template.users SET credits = credits - %s WHERE id = %s", (price, id))
    # Insert order record
    cursor.execute("""
        INSERT INTO template.Orders (user_id, case_type, template_type, price, payload)
        VALUES (%s, %s, %s, %s, %s)
    """, (id, caseType, templateType, price, str(replacements)))
    conn.commit()
    return True, "", 200


if __name__ == '__main__':
    updateRecordInTable("Test", [
        {'id': 1, 'caseTypeid': '2', 'templateName': 'Rental Lease Templates', 'isActive': 1, 'updated': True},
        {'id': 2, 'caseTypeid': '2', 'templateName': 'Rental Lease Templates', 'isActive': 1, 'added': True}
    ])
