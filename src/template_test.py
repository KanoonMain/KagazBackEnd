import re
from docx import Document
from docx2pdf import convert
import pikepdf
def replace_placeholders_in_paragraph(paragraph, replacements):
    # Combine runs into full text
    full_text = ''.join(run.text for run in paragraph.runs)

    # Replace placeholders like [LandlordFullName] with data values
    for key, value in replacements.items():
        pattern = rf"\[{re.escape(key)}\]"
        full_text = re.sub(pattern, value, full_text)

    # Clear all existing runs
    for run in paragraph.runs:
        run.text = ''

    # Set the new text to the first run
    if paragraph.runs:
        paragraph.runs[0].text = full_text
    else:
        paragraph.add_run(full_text)

def generateProtectedPDF(caseType, TemplateType, replacements):
    doc_path = "Rental_Agreement.docx"
    doc = Document(doc_path)
    temp_docx = "Filled_Rental_Agreement.docx"
    # Apply replacements to all paragraphs
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, replacements)
    doc.save(temp_docx)
    # === Step 4: Convert to PDF ===
    output_pdf = "Filled_Rental_Agreement.pdf"
    secured_pdf = "Secured_Rental_Agreement.pdf"
    convert(temp_docx, output_pdf)

    # === Step 5: Apply PDF security (disable copy & print) ===
    with pikepdf.open(output_pdf) as pdf:
        pdf.save(
            secured_pdf,
            encryption=pikepdf.Encryption(
                user="",
                owner="secret",  # Owner password
                allow=pikepdf.Permissions(
                        extract=False,
                        modify_annotation=False,
                        modify_assembly=False,
                        modify_form=False,
                        modify_other=False,
                        print_lowres=True,
                        print_highres=True)))


if __name__ == '__main__':
    data = {
        "LandlordFullName": "John Doe Landlord",
        "LandlordAddress": "Random Landlord",
        "Tenant Full Name": "John Doe Tenant ",
        "Tenant Address": "Random Tenant’s ",
        "Rental Property Address": "ABCD",
        "Start Date": "16-04-2025",
        "End Date": "16-04-2025",
        "Amount": "1000",
        "month/week": "Month",
        "day": "5th",
        "payment method": "NEFT",
        "SecurityAmount": "10000",
        "X": "10",
        "number": "30"
    }
    generateProtectedPDF("Civil Templates", "Rental Lease Templates", data)
