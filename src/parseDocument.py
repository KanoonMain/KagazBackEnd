import sqlite3
import re
from docx import Document
from connect_db import getConnection

# Step 1: Read content from Word file
doc_path = 'Rental_Agreement.docx'
doc = Document(doc_path)

# Extract full text from paragraphs and tables with spacing preserved
full_text = ''

# Read normal paragraphs
for para in doc.paragraphs:
    full_text += para.text + '\n'

# Read text inside tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            full_text += cell_text + '\n'

# Step 2: Extract placeholders in square brackets
placeholders = re.findall(r'\[([^\[\]]+)\]', full_text)
placeholdersData = {}
for item in placeholders:
    placeholdersData[item] = item
print(placeholdersData)

# Step 3: Read file as binary for blob storage
with open(doc_path, 'rb') as file:
    file_blob = file.read()

# Step 4: Insert into database
conn = getConnection()
cursor = conn.cursor()
cursor.execute(
    'INSERT INTO Templates (caseTypeid, TemplateTypeid, templateData, templateForm) VALUES (?, ?, ?, ?)',
    (1, 1, file_blob, str(placeholdersData))
)

# Finalize
conn.commit()
conn.close()